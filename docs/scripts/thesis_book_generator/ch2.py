# -*- coding: utf-8 -*-
"""Chapter 2: Literature Review and Research Gap."""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from build_thesis import (FONT, add_para, add_bullet, add_chapter_heading, add_section_heading,
                           add_table_with_caption, page_break)
from refs import cite


def _review_heading(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = FONT
    r.font.size = Pt(12.5)
    r.bold = True
    return p


def _system_review(doc, name, contribution, limitations, gap, citation=None):
    """Add one compact literature-review block with scannable categories."""
    title = f"{name} {citation}" if citation else name
    _review_heading(doc, title)
    add_bullet(doc, contribution, bold_lead="Contribution: ")
    add_bullet(doc, limitations, bold_lead="Limitations: ")
    add_bullet(doc, gap, bold_lead="Gap for AEGIS: ")


def chapter2(doc):
    add_chapter_heading(doc, 2, "Literature Review and Research Gap")

    add_para(doc,
              "The literature review focuses on the work most directly comparable to AEGIS. It is "
              "organized as short review blocks so each source can be checked quickly by contribution, "
              "limitation, and the specific gap it leaves for this thesis.", space_after=0)

    # ---------------------------------------------------------------- 2.1
    add_section_heading(doc, "2.1", "Natural Language Interfaces to Databases")
    _system_review(
        doc,
        "NLIDB surveys",
        "Compare major natural-language database interface designs, including keyword, pattern, parsing, and grammar-based systems.",
        "Safety is treated mostly as a secondary concern; even SQL-injection discussion ends with generic filtering advice.",
        "No reviewed system makes SQL safety a structural design property.",
        cite('affolter19', 'liu_xu25'))
    _system_review(
        doc,
        "NaLIR",
        "Uses an intermediate query tree and asks the user to choose between possible interpretations.",
        "Correctness depends on user disambiguation, and the final output remains executable SQL.",
        "Ambiguity is handled after parsing, but unsafe or unauthorized queries are not prevented by design.",
        cite('li_jagadish14'))
    _system_review(
        doc,
        "Veezoo",
        "Uses an editable Knowledge Graph that maps business language to database concepts.",
        "The semantic layer improves usability, but the paper does not evaluate SQL safety or permission enforcement.",
        "A semantic layer exists in prior work, but not as a safety boundary.",
        cite('lehmann22'))

    # ---------------------------------------------------------------- 2.2
    add_section_heading(doc, "2.2", "Neural and LLM-Based Text-to-SQL")
    _system_review(
        doc,
        "Spider and BIRD",
        "Provide large text-to-SQL benchmarks for complex and cross-domain database questions.",
        "They measure SQL correctness and execution accuracy, not whether a query should be allowed.",
        "Benchmark success does not equal database safety or business-policy compliance.",
        cite('yu_spider18', 'li_bird23'))
    _system_review(
        doc,
        "RAT-SQL",
        "Models schema relations directly so table and column linking can improve on complex schemas.",
        "Errors still frequently come from wrong table or column selection.",
        "Better schema encoding reduces some mistakes but does not create an authorization boundary.",
        cite('wang_rat20'))
    _system_review(
        doc,
        "PICARD",
        "Constrains decoding so generated SQL is more likely to be syntactically valid.",
        "The model still authors the SQL, and grammar validity does not prove the query is semantically safe.",
        "A valid SQL string can still express the wrong intent or access the wrong data.",
        cite('scholak21'))
    _system_review(
        doc,
        "G-SQL and TriSQL",
        "Show two strong directions: rule-guided schema-aware translation and multi-stage LLM refinement.",
        "G-SQL has limited qualitative evaluation, while TriSQL still depends on LLM-authored executable SQL.",
        "The closest prior work improves accuracy, but does not remove SQL generation from the LLM output space.",
        cite('shalaan25', 'su_trisql26'))

    # ---------------------------------------------------------------- 2.3
    add_section_heading(doc, "2.3", "Natural Language for Visualization and Dashboards")
    _system_review(
        doc,
        "nl4dv and DataTone",
        "Map natural-language questions into visualization tasks and expose ambiguity instead of silently guessing.",
        "They focus on visualization specification, not protected database execution or persistent dashboard widgets.",
        "Visualization systems solve chart selection but leave SQL safety and reuse outside the architecture.",
        cite('narechania21', 'gao15'))
    _system_review(
        doc,
        "DashBot",
        "Uses deep reinforcement learning and dashboard-design rules to compose multi-chart dashboards.",
        "It has no natural-language-to-SQL stage and does not handle governed production database access.",
        "Dashboard composition is useful, but it does not solve safe language-to-data translation.",
        cite('deng23'))

    # ---------------------------------------------------------------- 2.4
    add_section_heading(doc, "2.4", "Applied Conversational Business Intelligence")
    _system_review(
        doc,
        "Conversational BI assistants",
        "Demonstrate that LLM-based chat can be connected to dashboards and SQL tools.",
        "Direct SQL execution through an LLM tool loop creates a large attack surface and is rarely evaluated adversarially.",
        "A production assistant needs a constrained execution layer, not only a conversational interface.",
        cite('shailesh25'))
    _system_review(
        doc,
        "Explanatory dashboard analytics",
        "Shows that deterministic computation followed by LLM narration can improve reliability in business explanation tasks.",
        "The work targets dashboard explanation, not safe SQL compilation or reusable widget creation.",
        "It supports AEGIS's choice to let deterministic code compute results while the LLM handles language.",
        cite('valkenburgh24'))

    # ---------------------------------------------------------------- 2.5
    add_section_heading(doc, "2.5", "Comparative Summary")
    add_table_with_caption(
        doc, "Table 1: Comparative summary of the most closely related systems.",
        ["System", "NL", "Semantic", "Safe SQL", "Visual",
         "Widget", "Coverage", "Evaluation"],
        [
            ["Spider/BIRD", "Yes", "-", "-", "-", "-", "-", "Benchmark"],
            ["RAT-SQL/PICARD", "Yes", "-", "Partial", "-", "-", "-", "Benchmark"],
            ["G-SQL/TriSQL", "Yes", "Partial", "Partial", "-", "-", "-", "Benchmark"],
            ["NaLIR", "Yes", "-", "-", "-", "-", "-", "User study"],
            ["Veezoo", "Yes", "Yes", "-", "-", "-", "-", "User study"],
            ["nl4dv/DataTone", "Yes", "-", "-", "Yes", "-", "-", "User study"],
            ["DashBot", "-", "-", "-", "Yes", "Partial", "-", "Synthetic"],
            ["Conversational BI", "Yes", "-", "-", "Yes", "-", "-", "Demo"],
            ["AEGIS", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "Production"],
        ],
        col_widths=[1.25, 0.55, 0.75, 0.65, 0.65, 0.65, 0.75, 1.10],
        font_size=9.0,
        keep_together=True,
        caption_space_before=4)
    page_break(doc)

    # ---------------------------------------------------------------- 2.6
    add_section_heading(doc, "2.6", "Research Gap Analysis")
    add_bullet(doc, "Most systems measure answer accuracy, but do not evaluate unsafe SQL behavior.", bold_lead="Safety gap: ")
    add_bullet(doc, "Semantic layers appear in prior work, but mainly for usability and matching, not as execution boundaries.", bold_lead="Semantic-layer gap: ")
    add_bullet(doc, "Visualization systems produce charts, but usually operate outside governed SQL execution.", bold_lead="Visualization gap: ")
    add_bullet(doc, "Prior tools often answer one question at a time, while recurring business reports need refreshable outputs.", bold_lead="Persistence gap: ")
    add_bullet(doc, "AEGIS addresses these gaps through a bounded vocabulary, deterministic query compiler, safe visualization selector, and reusable widget model.", bold_lead="Thesis position: ")
    page_break(doc)


