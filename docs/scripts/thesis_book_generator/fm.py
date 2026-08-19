# -*- coding: utf-8 -*-
"""Front matter: title pages, certifications, acknowledgement, TOC, LOF, LOT."""
from pathlib import Path

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from build_thesis import add_para, add_mixed_para, page_break, add_tab_leader, FONT

TITLE = "AEGIS: A Constraint-Based Architecture for Safe LLM-Assisted Natural Language Analytics"
SUPERVISOR = "Mst. Sahela Rahman"
STUDENT = "Md. Riaz"
STUDENT_ID = "0322310105101024"
SUBMISSION_DATE = "Friday, August 07, 2026"
LOGO_PATH = Path(__file__).with_name("pundra_logo.png")


def _title_block(doc, eyebrow, include_logo=True):
    """Shared title text used by the cover and signature title page."""
    if include_logo and LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(12)
        p_logo.add_run().add_picture(str(LOGO_PATH), width=Inches(1.15))
    add_para(doc, eyebrow, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18, line_spacing=1.2)
    add_para(doc, f'“{TITLE}”', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=20, line_spacing=1.18)
    add_para(doc, "This thesis is submitted to the department of Computer Science & Engineering in "
                  "partial fulfillment of the requirements for the degree of",
              size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.2)
    add_para(doc, "Bachelor of Science in Computer Science & Engineering", size=12, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=32, line_spacing=1.2)
    add_para(doc, "Under the Course of-", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28, line_spacing=1.2)
    add_para(doc, "Course Title: Thesis/Project Work", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, "Course Code: CSE 4000(B)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=34, line_spacing=1.1)


def _signature_title_block(doc):
    """Compact repeat title for the signature page."""
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(12)
        p_logo.add_run().add_picture(str(LOGO_PATH), width=Inches(1.15))
    add_para(doc, "A Thesis on-", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18, line_spacing=1.2)
    add_para(doc, f'“{TITLE}”', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=28, line_spacing=1.18)
    add_para(doc, "This thesis is submitted to the department of Computer Science & Engineering in "
                  "partial fulfillment of the requirements for the degree of",
              size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.2)
    add_para(doc, "Bachelor of Science in Computer Science & Engineering", size=12, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=88, line_spacing=1.2)


def title_page(doc):
    _title_block(doc, "A Thesis on-")
    add_para(doc, "Thesis Supervisor-", size=12.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, f"Name: {SUPERVISOR}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, "Designation: Lecturer", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=34, line_spacing=1.1)
    add_para(doc, "Prepared by-", size=12.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, f"Name: {STUDENT}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, f"ID/Registration No: {STUDENT_ID}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, "Batch: 16th", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, "Semester: 8th", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, "Session: Spring - 2023", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30, line_spacing=1.1)
    add_para(doc, "DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, line_spacing=1.1)
    add_para(doc, "PUNDRA UNIVERSITY OF SCIENCE & TECHNOLOGY", size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22, line_spacing=1.1)
    add_para(doc, "Date of submission: ____________________", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.1)
    page_break(doc)


def signature_title_page(doc):
    _signature_title_block(doc)
    add_para(doc, "----------------------------------------", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.1)
    add_para(doc, "Signature of Supervisor", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=56, line_spacing=1.1)
    add_para(doc, "----------------------------------------", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.1)
    add_para(doc, "Signature of Student", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=82, line_spacing=1.1)
    add_para(doc, "DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, line_spacing=1.1)
    add_para(doc, "PUNDRA UNIVERSITY OF SCIENCE & TECHNOLOGY", size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20, line_spacing=1.1)
    add_para(doc, "Date of submission: ____________________", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.1)
    page_break(doc)


def certification_of_originality(doc):
    add_para(doc, "CERTIFICATION OF ORIGINALITY", size=15, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
    add_para(doc,
             f'I hereby declare that this thesis titled "{TITLE}" is my own original work, carried out '
             f'under the supervision of {SUPERVISOR}, Lecturer, Department of Computer Science & '
             "Engineering, Pundra University of Science & Technology. To the best of my knowledge, this "
             "thesis contains no material previously published or written by another person, except where "
             "due reference is made in the text. This work has not been submitted, in whole or in part, for "
             "any other degree or qualification at this or any other institution.",
             size=12, space_after=40)
    add_para(doc, "____________________________", space_after=2)
    add_para(doc, "Signature of Student", size=12, space_after=2)
    add_para(doc, STUDENT, size=12, bold=True, space_after=2)
    add_para(doc, f"ID: {STUDENT_ID}", size=12, space_after=0)
    page_break(doc)


def certification_of_approval(doc):
    add_para(doc, "CERTIFICATION OF APPROVAL", size=15, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc,
             f'This is to certify that the research paper titled "{TITLE}"',
             size=12, space_after=16, line_spacing=1.5)
    add_para(doc,
             "That it has been read, evaluated and accepted for fulfilment of the Academic Degree of "
             "Bachelor in Computer Science and Engineering (B.S.C. in CSE) at Pundra University of "
             "Science & Technology.",
             size=12, space_after=16, line_spacing=1.5)
    add_para(doc,
             "The acceptance decision is made based upon an independent review process that provides "
             "critically constructive feedback within a short turnaround time. This paper represents "
             "the author's independent work, critical analysis, and capability in undertaking "
             "literature research as per the academic policies and ethical standards of the university.",
             size=12, space_after=16, line_spacing=1.5)
    add_para(doc,
             "I certify that the candidate has completed the required research activities for the "
             "program and recommend that this piece of work as of acceptable standard for submission "
             "and for inclusion in the academic record of the University.",
             size=12, space_after=28, line_spacing=1.5)

    add_para(doc, SUPERVISOR, size=12, bold=True, space_after=0, line_spacing=1.0)
    add_para(doc, "Lecturer", size=12, space_after=0, line_spacing=1.0)
    add_para(doc, "Department of Computer Science and Engineering", size=12, space_after=0, line_spacing=1.0)
    add_para(doc, "Pundra University of Science & Technology", size=12, space_after=48, line_spacing=1.0)

    add_mixed_para(doc, [("Date: ", True, False), ("____________________________", False, False)],
                   space_after=0, line_spacing=1.0)
    page_break(doc)


def acknowledgement(doc):
    add_para(doc, "ACKNOWLEDGEMENT", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc,
             "I would like to express my sincere gratitude to my supervisor, Mst. Sahela Rahman, Lecturer, "
             "Department of Computer Science & Engineering, Pundra University of Science & Technology, for "
             "her continuous guidance, patience, and constructive feedback throughout the design, "
             "implementation, and evaluation of this research. Her insistence on precise, defensible claims "
             "shaped this thesis at every stage, from the formal safety proof to the honesty of its "
             "discussion of limitations.",
             size=12, space_after=14)
    add_para(doc,
             "I am also grateful to the faculty members of the Department of Computer Science & Engineering "
             "for the coursework and feedback that prepared me to undertake this research, and to my family "
             "for their patience and encouragement during the long implementation and evaluation cycles this "
             "thesis required.",
             size=12, space_after=14)
    add_para(doc,
             "Finally, I acknowledge the authors whose published research on natural language "
             "interfaces, text-to-SQL translation, and dashboard generation provided the foundation "
             "against which AEGIS's contributions are measured.",
             size=12, space_after=0)
    page_break(doc)


TOC_ENTRIES = [
    (0, "Certification of Originality", "i"),
    (0, "Certification of Approval", "ii"),
    (0, "Acknowledgement", "iii"),
    (0, "Abstract", "iv"),
    (0, "Chapter 1: Introduction", "1"),
    (1, "1.1 Background", "1"),
    (1, "1.2 Problem Statement", "1"),
    (1, "1.3 Research Novelty and Motivation", "2"),
    (1, "1.4 Objectives and Contributions", "3"),
    (0, "Chapter 2: Literature Review and Research Gap", "4"),
    (1, "2.1 Natural Language Interfaces to Databases", "4"),
    (1, "2.2 Neural and LLM-Based Text-to-SQL", "4"),
    (1, "2.3 Natural Language for Visualization and Dashboards", "5"),
    (1, "2.4 Applied Conversational Business Intelligence", "6"),
    (1, "2.5 Comparative Summary", "7"),
    (1, "2.6 Research Gap Analysis", "7"),
    (0, "Chapter 3: Methodology", "8"),
    (1, "3.1 Research Paradigm", "8"),
    (1, "3.2 Formative Study of Reporting Patterns", "8"),
    (1, "3.3 Design Principles", "9"),
    (1, "3.4 Formal Model", "10"),
    (1, "3.5 Threat Model", "10"),
    (1, "3.6 System Architecture", "12"),
    (1, "3.7 Semantic Layer Design", "13"),
    (1, "3.8 Intent Parsing with Dynamic Vocabulary Injection", "15"),
    (1, "3.9 Safe Query Compiler", "16"),
    (1, "3.10 Visualization Selector", "18"),
    (1, "3.11 Widget Persistence and Reuse", "19"),
    (0, "Chapter 4: Experimental Work", "20"),
    (1, "4.1 Implementation", "20"),
    (1, "4.2 Experimental Environment", "21"),
    (1, "4.3 Benchmark Dataset Construction", "21"),
    (1, "4.4 Baseline and Oracle Comparisons", "22"),
    (1, "4.5 Evaluation Procedure", "23"),
    (0, "Chapter 5: Results and Discussion", "24"),
    (1, "5.1 Evaluation Overview", "24"),
    (1, "5.2 500-Question Natural-Language Benchmark", "24"),
    (1, "5.3 Admin Fidelity Benchmark", "25"),
    (1, "5.4 Focused Semantic Coverage", "25"),
    (1, "5.5 Safety Evaluation", "26"),
    (1, "5.6 Comparison With Direct LLM-to-SQL", "26"),
    (1, "5.7 Interpretation", "27"),
    (0, "Chapter 6: Limitations and Future Work", "28"),
    (1, "6.1 Limitations", "28"),
    (1, "6.2 Future Work", "29"),
    (0, "Chapter 7: Conclusion", "30"),
    (0, "References", "31"),
]


def table_of_contents(doc):
    add_para(doc, "Table of Contents", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    for level, text, pg in TOC_ENTRIES:
        p = add_para(doc, space_after=4 if level else 8, space_before=0 if level else 4)
        indent = 0.3 if level else 0.0
        p.paragraph_format.left_indent = Inches(indent)
        add_tab_leader(p, 6.0, leader='dot')
        r = p.add_run(text + '\t' + pg)
        r.font.name = FONT
        r.font.size = Pt(12 if level == 0 else 11.5)
        r.bold = (level == 0)
    page_break(doc)


LOF = [
    ("Figure 1: Design Science Research workflow", "8"),
    ("Figure 2: AEGIS architecture pipeline", "13"),
    ("Figure 3: Semantic layer modularity", "14"),
    ("Figure 4: Vocabulary injection and structured intent validation workflow", "16"),
    ("Figure 5: Two-layer SQL safety defence", "18"),
    ("Figure 6: Widget lifecycle and refresh model", "19"),
]

LOT = [
    ("Table 2.1: Comparative summary of related systems", "7"),
    ("Table 3.1: AEGIS reporting pattern taxonomy", "9"),
    ("Table 3.2: Semantic layer implementation contract", "14"),
    ("Table 3.3: AEGIS analytical patterns", "17"),
    ("Table 3.4: Visualization selector mapping", "19"),
    ("Table 4.1: Prototype module-to-architecture mapping", "21"),
    ("Table 4.2: Experimental setup", "21"),
    ("Table 4.3: Static evaluation datasets", "22"),
    ("Table 5.1: Current evaluation artifacts", "24"),
    ("Table 5.2: 500-question live benchmark results", "25"),
    ("Table 5.3: Admin analytics oracle benchmark", "25"),
    ("Table 5.4: Focused semantic coverage results", "26"),
    ("Table 5.5: Safety interpretation", "26"),
    ("Table 5.6: Structural comparison", "27"),
]


def list_of_figures(doc):
    add_para(doc, "List of Figures", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    for text, pg in LOF:
        p = add_para(doc, space_after=4)
        add_tab_leader(p, 6.0, leader='dot')
        r = p.add_run(text + '\t' + pg)
        r.font.name = FONT
        r.font.size = Pt(12)
    page_break(doc)


def list_of_tables(doc):
    add_para(doc, "List of Tables", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    for text, pg in LOT:
        p = add_para(doc, space_after=4)
        add_tab_leader(p, 6.0, leader='dot')
        r = p.add_run(text + '\t' + pg)
        r.font.name = FONT
        r.font.size = Pt(12)
    page_break(doc)
