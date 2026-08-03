# -*- coding: utf-8 -*-
"""AEGIS Thesis Book generator - rebuilt from AEGIS_Manuscript.md + EXPLAINER.md + references/*.pdf analysis."""
import copy
import os
import tempfile
import zipfile
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'Times New Roman'
BODY_SIZE = Pt(12)

# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def set_cell_text(cell, text, bold=False, size=11, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    if align:
        p.alignment = align
    return p


def style_table(table, header_rows=1):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(10.5)
                if ri < header_rows:
                    for r in p.runs:
                        r.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '999999')
        borders.append(el)
    tblPr.append(borders)


def set_table_keep_together(table):
    """Keep table rows from breaking across pages in Word/PDF export."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = trPr.find(qn('w:cantSplit'))
        if cant_split is None:
            cant_split = OxmlElement('w:cantSplit')
            trPr.append(cant_split)
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                keep = pPr.find(qn('w:keepNext'))
                if keep is None:
                    keep = OxmlElement('w:keepNext')
                    pPr.append(keep)


def set_page_number_format(section, fmt, start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    run.font.name = FONT
    run.font.size = Pt(11)


def add_footer_page_number(section, align=WD_ALIGN_PARAGRAPH.CENTER):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ''
    p.alignment = align
    add_page_number_field(p)
    return p


def set_hanging_indent(paragraph, left_in, hang_in):
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), str(int(left_in * 1440)))
    ind.set(qn('w:hanging'), str(int(hang_in * 1440)))


def add_tab_leader(paragraph, position_in, leader='dot'):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), leader)
    tab.set(qn('w:pos'), str(int(position_in * 1440)))
    tabs.append(tab)
    pPr.append(tabs)


# ---------------------------------------------------------------------------
# Document-level setup
# ---------------------------------------------------------------------------

def new_document():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    return doc


def new_section(doc, page_num_fmt, start=1):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    set_page_number_format(section, page_num_fmt, start)
    add_footer_page_number(section)
    return section


def add_para(doc, text='', size=12, bold=False, italic=False, align=None,
             space_after=6, space_before=0, indent=False, line_spacing=1.5,
             color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        pf.first_line_indent = Inches(0.4)
    if text:
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def add_mixed_para(doc, runs, align=None, space_after=6, space_before=0,
                    line_spacing=1.5):
    """runs: list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    for item in runs:
        text, bold, italic = (item + (False, False))[:3] if len(item) < 3 else item
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(12)
        r.bold = bold
        r.italic = italic
    return p


def add_chapter_heading(doc, chapter_no, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    r = p.add_run(f'CHAPTER {chapter_no}')
    r.font.name = FONT
    r.font.size = Pt(16)
    r.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    r2 = p2.add_run(title.upper())
    r2.font.name = FONT
    r2.font.size = Pt(16)
    r2.bold = True
    return p2


def add_section_heading(doc, number, title, level=2):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16 if level == 2 else 10)
    pf.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f'{number}  {title}')
    r.font.name = FONT
    r.bold = True
    r.font.size = Pt(14) if level == 2 else Pt(12.5)
    if level == 3:
        r.italic = True
    return p


def add_bullet(doc, text, level=0, bold_lead=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.5
    pf.left_indent = Inches(0.32 + (level * 0.25))
    pf.first_line_indent = Inches(-0.18)
    marker = p.add_run("\u2022\t")
    marker.font.name = FONT
    marker.font.size = Pt(12)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.name = FONT
        r1.font.size = Pt(12)
        r2 = p.add_run(text)
        r2.font.name = FONT
        r2.font.size = Pt(12)
    else:
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(12)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(12)
    return p


def add_table_with_caption(doc, caption, headers, rows, col_widths=None, font_size=10.0,
                           keep_together=True, caption_space_before=10):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(caption_space_before)
    cap.paragraph_format.space_after = Pt(4)
    if keep_together:
        cap._p.get_or_add_pPr().append(OxmlElement('w:keepNext'))
    rc = cap.add_run(caption)
    rc.bold = True
    rc.font.name = FONT
    rc.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    if col_widths:
        table.autofit = False
        for j, width in enumerate(col_widths):
            table.columns[j].width = Inches(width)
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, size=font_size,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i + 1].cells[j], str(val), size=font_size,
                          align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT)
            if col_widths:
                table.rows[i + 1].cells[j].width = Inches(col_widths[j])
    style_table(table)
    if keep_together:
        set_table_keep_together(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.line_spacing = 1.15
    for i, line in enumerate(code_text.strip('\n').split('\n')):
        r = p.add_run(line if i == 0 else '\n' + line)
        r.font.name = FONT
        r.font.size = Pt(9.5)
    return p


def enforce_docx_font(docx_path, font_name=FONT):
    """Normalize saved DOCX font declarations to the thesis font.

    python-docx leaves several built-in Word styles and list definitions in the
    package even when they are not explicitly authored. Normalizing the OOXML
    after save prevents stray Calibri/Courier/Symbol declarations from leaking
    into the final file.
    """
    out_dir = os.path.dirname(os.path.abspath(docx_path)) or "."
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=out_dir) as tmp:
        tmp_path = tmp.name
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    text = data.decode("utf-8", errors="ignore")
                    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                        text = __import__("re").sub(
                            rf'w:{attr}="[^"]+"',
                            f'w:{attr}="{font_name}"',
                            text,
                        )
                    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
                        text = __import__("re").sub(rf'\s+w:{attr}="[^"]+"', "", text)
                    data = text.encode("utf-8")
                zout.writestr(item, data)
        os.replace(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def add_figure_placeholder(doc, fig_num, title, description, height_in=2.6):
    """Dashed-border placeholder box marking where a real diagram/chart belongs.

    Once the actual image file exists, replace the call site with
    doc.add_picture(path, width=Inches(6.0)) followed by the same caption.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.3)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'dashed')
        el.set(qn('w:sz'), '10')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '888888')
        borders.append(el)
    tcPr.append(borders)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F7F7F7')
    tcPr.append(shd)

    cell.text = ''
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(14)
    p1.paragraph_format.space_after = Pt(6)
    fig_label = str(fig_num)
    r1 = p1.add_run(f'[ PLACEHOLDER - FIGURE {fig_label} NOT YET INSERTED ]')
    r1.bold = True
    r1.font.name = FONT
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x99, 0x33, 0x00)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(title)
    r2.bold = True
    r2.font.name = FONT
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    row = table.rows[0]
    row.height = Inches(height_in)
    from docx.oxml.ns import qn as _qn
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(_qn('w:val'), str(int(height_in * 1440)))
    trHeight.set(_qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(14)
    rc = cap.add_run(f'Figure {fig_label}: {title}')
    rc.bold = True
    rc.font.name = FONT
    rc.font.size = Pt(10.5)
    return table
